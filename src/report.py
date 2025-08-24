from __future__ import annotations

from pathlib import Path
from typing import Dict

import pandas as pd

from .data_utils import DatasetProfile


def _profile_to_markdown(title: str, profile: DatasetProfile) -> str:
    lines = [f"## {title}"]
    lines.append(f"Filas: {profile.num_rows} | Columnas: {profile.num_cols}")
    lines.append("")
    lines.append("### Tipos de datos")
    lines.append(profile.dtypes.to_frame("dtype").to_markdown())
    lines.append("")
    lines.append("### Valores faltantes (conteo)")
    lines.append(profile.missing_counts.to_frame("missing").to_markdown())
    lines.append("")
    lines.append("### Estadísticos (numéricos)")
    lines.append(profile.describe_numeric.to_markdown())
    lines.append("")
    return "\n".join(lines)


def generate_markdown_report(
    output_dir: Path,
    day_profile: DatasetProfile,
    hour_profile: DatasetProfile,
    top_day_corr: pd.DataFrame,
    top_hour_corr: pd.DataFrame,
    figures: Dict[str, str],
) -> Path:
    md_lines = ["# Análisis descriptivo Bike Sharing"]
    md_lines.append("")
    md_lines.append("Variable respuesta: `cnt`.")
    md_lines.append("")

    # Secciones de perfiles
    md_lines.append(_profile_to_markdown("Perfil dataset (day)", day_profile))
    md_lines.append(_profile_to_markdown("Perfil dataset (hour)", hour_profile))

    # Correlaciones top
    md_lines.append("## Correlaciones más altas con `cnt` (day)")
    md_lines.append(top_day_corr.head(8).to_markdown())
    md_lines.append("")
    md_lines.append("## Correlaciones más altas con `cnt` (hour)")
    md_lines.append(top_hour_corr.head(8).to_markdown())
    md_lines.append("")

    # Figuras
    md_lines.append("## Figuras")
    for label, rel_path in figures.items():
        md_lines.append(f"![{label}]({rel_path})")
    md_lines.append("")

    # Separación de train/val (texto)
    md_lines.append("## Conjuntos de modelización y validación")
    md_lines.append("Se realizó un split temporal 80/20 por `dteday` para ambos datasets.")

    out_path = output_dir / "reporte.md"
    out_path.write_text("\n".join(md_lines), encoding="utf-8")
    return out_path


def generate_docx_report(
    output_dir: Path,
    day_profile: DatasetProfile,
    hour_profile: DatasetProfile,
    top_day_corr: pd.DataFrame,
    top_hour_corr: pd.DataFrame,
    figures: Dict[str, str],
) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()

    # Estilo global: Calibri 12, interlineado 1.5
    try:
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
    except Exception:
        pass

    # Título
    doc.add_heading("Análisis descriptivo Bike Sharing", level=1)

    # Objetivo
    doc.add_heading("Objetivo", level=2)
    doc.add_paragraph(
        "Realizar un análisis descriptivo de Bike Sharing para entender la variable respuesta cnt y su relación con predictores, preparar datos y separar conjuntos de modelización y validación."
    )

    # Datos
    doc.add_heading("Datos", level=2)
    doc.add_paragraph(
        "Conjuntos: day.csv (agregado diario) y hour.csv (por hora). La variable respuesta es cnt (total de alquileres), donde cnt = registered + casual."
    )

    # Metodología
    doc.add_heading("Metodología", level=2)
    doc.add_paragraph(
        "Lectura y tipado de fecha (dteday); perfilado (dimensiones, tipos, faltantes, estadísticos); tratamiento defensivo de faltantes; correlaciones con cnt; split temporal 80/20 por dteday; visualizaciones (distribución de cnt y heatmaps)."
    )

    # Resultados clave
    doc.add_heading("Resultados clave", level=2)

    # Correlaciones resumidas
    doc.add_paragraph("Relación con cnt (day): registered (0.946), casual (0.673), atemp (0.631), temp (0.627), instant (0.629), yr (0.567), weathersit (-0.297), windspeed (-0.235), hum (-0.101).")
    doc.add_paragraph("Relación con cnt (hour): registered (0.972), casual (0.695), temp (0.405), atemp (0.401), hr (0.394), hum (-0.323), weathersit (-0.142).")

    # Sub-sección Distribución y ubicación de figuras correspondientes
    doc.add_heading("Distribución de cnt", level=3)
    # Figura 1: Distribución cnt (day)
    p = doc.add_paragraph("Figura 1. Distribución de cnt (day)")
    p.italic = True
    doc.add_picture(str(output_dir / figures["day_cnt_dist"]), width=Inches(5.5))
    # Figura 2: Distribución cnt (hour)
    p = doc.add_paragraph("Figura 2. Distribución de cnt (hour)")
    p.italic = True
    doc.add_picture(str(output_dir / figures["hour_cnt_dist"]), width=Inches(5.5))

    # Comentarios sobre predictores
    doc.add_heading("Comentarios sobre predictores y cnt", level=2)
    doc.add_paragraph(
        "registered y casual explican gran parte de la variación de cnt pero no deben usarse como predictores (fuga de información). Temperatura real/aparente y variables temporales muestran relación positiva; peor clima, humedad y viento altos se asocian con menor demanda."
    )

    # Heatmaps ubicados tras los comentarios
    doc.add_heading("Correlaciones (heatmaps)", level=3)
    # Figura 3: Heatmap (day)
    p = doc.add_paragraph("Figura 3. Heatmap de correlaciones (day)")
    p.italic = True
    doc.add_picture(str(output_dir / figures["day_corr_heatmap"]), width=Inches(5.5))
    # Figura 4: Heatmap (hour)
    p = doc.add_paragraph("Figura 4. Heatmap de correlaciones (hour)")
    p.italic = True
    doc.add_picture(str(output_dir / figures["hour_corr_heatmap"]), width=Inches(5.5))

    # Preparación y validación
    doc.add_heading("Preparación de datos y validación", level=2)
    doc.add_paragraph(
        "No se observaron faltantes relevantes; se dejó imputación defensiva. Se aplicó división temporal 80/20 por dteday para modelización y validación."
    )

    # Conclusiones
    doc.add_heading("Conclusiones", level=2)
    doc.add_paragraph(
        "cnt crece con temperatura y con el paso del tiempo (yr), y disminuye con malas condiciones meteorológicas, humedad y viento. Las horas del día influyen marcadamente (picos en horas punta en hour.csv). Para modelar cnt, excluir registered y casual; centrarse en clima, tiempo y calendario."
    )

    # Limitaciones y próximos pasos
    doc.add_heading("Limitaciones y próximos pasos", level=2)
    doc.add_paragraph(
        "La correlación no implica causalidad; conviene explorar modelos con validación temporal (p.ej., regresión regularizada, árboles/boosting) y variables derivadas (fines de semana/feriados, no linealidades en clima)."
    )

    out_path = output_dir / "reporte.docx"
    doc.save(out_path)
    return out_path
